# -*- coding: utf-8 -*-

import os
import io
from flask import Flask, request, send_file, render_template, jsonify
from docx import Document
import pandas as pd
from zipfile import ZipFile
from datetime import datetime
import logging # Importa o módulo de logging

# --- Configuração Inicial ---
app = Flask(__name__, template_folder='.')
PASTA_SAIDA_TEMP = "temp_output"
os.makedirs(PASTA_SAIDA_TEMP, exist_ok=True)

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Caminhos para os Modelos ---
TEMPLATE_GQ = "Portariagq.docx"
TEMPLATE_REMOCAO = {
    "oficio_com_ajuda": "remocao_oficio_com_ajuda.docx",
    "oficio_sem_ajuda": "remocao_oficio_sem_ajuda.docx",
    "a_pedido": "remocao_a_pedido.docx",
    "a_pedido_conjuge": "remocao_a_pedido_conjuge.docx"
}
TEMPLATE_VACANCIA = {
    "a_pedido": "vacancia_a_pedido.docx",
    "inacumulavel": "vacancia_inacumulavel.docx"
}
TEMPLATE_GSISTE = {
    "concessao_622": "gsiste_concessao_622.docx",
    "concessao_654": "gsiste_concessao_654.docx",
    "exclusao_622": "gsiste_exclusao_622.docx",
    "exclusao_654": "gsiste_exclusao_654.docx"
}


# --- Funções Auxiliares ---
def get_value(dados, key_options):
    for key in key_options:
        val = dados.get(key)
        if val is not None and not pd.isnull(val): return val
    return None

def formatar_data_ddmmaaaa(data_obj):
    if pd.isnull(data_obj) or data_obj == '': return ""
    try: return pd.to_datetime(data_obj).strftime("%d/%m/%Y")
    except (ValueError, TypeError): return str(data_obj)

def preencher_documento(doc, substituicoes):
    for p in doc.paragraphs:
        for key, value in substituicoes.items(): p.text = p.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in substituicoes.items(): p.text = p.text.replace(key, str(value))
    return doc

# --- Funções de Geração (sem alterações) ---

def criar_portaria_gq(dados: dict):
    try:
        doc = Document(TEMPLATE_GQ)
        logging.info(f"Modelo GQ '{TEMPLATE_GQ}' carregado com sucesso.")
    except Exception as e:
        logging.error(f"Erro ao carregar o modelo GQ '{TEMPLATE_GQ}': {e}")
        return None, None
    gq_map = {"GQI": "GQ1", "GQII": "GQ2", "GQIII": "GQ3"}
    nome_servidor = get_value(dados, ['NOME DO SERVIDOR', 'SERVIDOR']) or 'ServidorDesconhecido'
    substituicoes = {
        "#PROCESSO": str(get_value(dados, ['PROCESSO']) or ''), "#GQ": gq_map.get(str(get_value(dados, ['TIPO DE GQ', 'GQ']) or '').strip().upper(), ''),
        "#SERVIDOR": str(nome_servidor), "#CPF": str(get_value(dados, ['CPF']) or ''),
        "#SIAPE": str(get_value(dados, ['SIAPE']) or ''), "#DATAGQ": formatar_data_ddmmaaaa(get_value(dados, ['DATA DA GQ', 'DATAGQ'])),
    }
    doc = preencher_documento(doc, substituicoes)
    nome_ficheiro = f"Portaria_GQ_{nome_servidor.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    return nome_ficheiro, doc

def criar_portaria_remocao(dados: dict):
    tipo_remocao = dados.get('TIPO_REMOCAO')
    template_path = TEMPLATE_REMOCAO.get(tipo_remocao)
    if not template_path:
        logging.warning(f"Tipo de remoção '{tipo_remocao}' não encontrado nos modelos.")
        return None, None
    try:
        doc = Document(template_path)
        logging.info(f"Modelo de remoção '{template_path}' carregado com sucesso.")
    except Exception as e:
        logging.error(f"Erro ao carregar o modelo de remoção '{template_path}': {e}")
        return None, None
    nome_servidor = get_value(dados, ['SERVIDOR', 'NOME DO SERVIDOR']) or 'ServidorDesconhecido'
    substituicoes = {
        "#PROCESSO": str(get_value(dados, ['PROCESSO']) or ''), "#SERVIDOR": str(nome_servidor),
        "#CPF": str(get_value(dados, ['CPF']) or ''), "#SIAPE": str(get_value(dados, ['SIAPE']) or ''),
        "#CARGO": str(get_value(dados, ['CARGO']) or ''), "#LOTACAOORIGEM": str(get_value(dados, ['LOTACAOORIGEM']) or ''),
        "#LOTACAODESTINO": str(get_value(dados, ['LOTACAODESTINO']) or '')
    }
    data_vigencia = get_value(dados, ['DATA_VIGENCIA'])
    substituicoes["#CLAUSULA_VIGENCIA"] = f", a partir de {formatar_data_ddmmaaaa(data_vigencia)}." if data_vigencia else "."
    doc = preencher_documento(doc, substituicoes)
    nome_ficheiro = f"Portaria_Remocao_{nome_servidor.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    return nome_ficheiro, doc

def criar_portaria_vacancia(dados: dict):
    tipo_vacancia = dados.get('TIPO_VACANCIA')
    template_path = TEMPLATE_VACANCIA.get(tipo_vacancia)
    if not template_path:
        logging.warning(f"Tipo de vacância '{tipo_vacancia}' não encontrado nos modelos.")
        return None, None
    try:
        doc = Document(template_path)
        logging.info(f"Modelo de vacância '{template_path}' carregado com sucesso.")
    except Exception as e:
        logging.error(f"Erro ao carregar o modelo de vacância '{template_path}': {e}")
        return None, None
    nome_servidor = get_value(dados, ['NOME', 'SERVIDOR']) or 'ServidorDesconhecido'
    substituicoes = {
        "#PROCESSO": str(get_value(dados, ['PROCESSO']) or ''), "#CARGO": str(get_value(dados, ['CARGO']) or ''),
        "#CLASSE": str(get_value(dados, ['CLASSE']) or ''), "#PADRAO": str(get_value(dados, ['PADRAO']) or ''),
        "#NOME": str(nome_servidor), "#CPF": str(get_value(dados, ['CPF']) or ''),
        "#SIAPE": str(get_value(dados, ['SIAPE']) or ''), "#VACANCIA": formatar_data_ddmmaaaa(get_value(dados, ['VACANCIA'])),
        "#NOVOCARG": str(get_value(dados, ['NOVOCARG']) or ''), "#NOVOORG": str(get_value(dados, ['NOVOORG']) or '')
    }
    doc = preencher_documento(doc, substituicoes)
    nome_ficheiro = f"Portaria_Vacancia_{nome_servidor.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    return nome_ficheiro, doc

def criar_portaria_gsiste(dados: dict):
    tipo_gsiste = dados.get('TIPO_GSISTE')
    template_path = TEMPLATE_GSISTE.get(tipo_gsiste)
    if not template_path:
        logging.warning(f"Tipo GSISTE '{tipo_gsiste}' não encontrado nos modelos.")
        return None, None
    try:
        doc = Document(template_path)
        logging.info(f"Modelo GSISTE '{template_path}' carregado com sucesso.")
    except Exception as e:
        logging.error(f"Erro ao carregar o modelo GSISTE '{template_path}': {e}")
        return None, None
    nome_servidor = get_value(dados, ['SERVIDOR']) or 'ServidorDesconhecido'
    substituicoes = {
        "#PROCESSO": str(get_value(dados, ['PROCESSO']) or ''),
        "#SERVIDOR": str(nome_servidor),
        "#SIAPE": str(get_value(dados, ['SIAPE']) or ''),
        "#CARGO": str(get_value(dados, ['CARGO']) or ''),
        "#LOTACAO": str(get_value(dados, ['LOTACAO']) or ''),
    }
    if 'exclusao' in tipo_gsiste:
        data_exclusao = get_value(dados, ['EXCLUSAO'])
        substituicoes['#EXCLUSAO'] = formatar_data_ddmmaaaa(data_exclusao)
    doc = preencher_documento(doc, substituicoes)
    nome_ficheiro = f"Portaria_GSISTE_{nome_servidor.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    return nome_ficheiro, doc


# --- Funções de Validação (sem alterações) ---
def validar_colunas_gq(df_columns):
    mapeamento = {'PROCESSO': ['PROCESSO'], 'TIPO DE GQ': ['TIPO DE GQ', 'GQ'], 'NOME DO SERVIDOR': ['NOME DO SERVIDOR', 'SERVIDOR'], 'CPF': ['CPF'], 'SIAPE': ['SIAPE'], 'DATA DA GQ': ['DATA DA GQ', 'DATAGQ']}
    return [f"'{c}'" for c, a in mapeamento.items() if not any(n in df_columns for n in a)]
def validar_colunas_remocao(df_columns):
    mapeamento = {'PROCESSO': ['PROCESSO'], 'SERVIDOR': ['SERVIDOR', 'NOME DO SERVIDOR'], 'CPF': ['CPF'], 'SIAPE': ['SIAPE'], 'CARGO': ['CARGO'], 'LOTACAOORIGEM': ['LOTACAOORIGEM'], 'LOTACAODESTINO': ['LOTACAODESTINO']}
    return [f"'{c}'" for c, a in mapeamento.items() if not any(n in df_columns for n in a)]
def validar_colunas_vacancia(df_columns, tipo_vacancia):
    mapeamento_base = {'PROCESSO': ['PROCESSO'], 'CARGO': ['CARGO'], 'CLASSE': ['CLASSE'], 'PADRAO': ['PADRAO'], 'NOME': ['NOME', 'SERVIDOR'], 'CPF': ['CPF'], 'SIAPE': ['SIAPE'], 'VACANCIA': ['VACANCIA']}
    if tipo_vacancia == 'inacumulavel':
        mapeamento_base['NOVOCARG'] = ['NOVOCARG']; mapeamento_base['NOVOORG'] = ['NOVOORG']
    return [f"'{c}'" for c, a in mapeamento_base.items() if not any(n in df_columns for n in a)]
def validar_colunas_gsiste(df_columns, tipo_gsiste):
    mapeamento_base = {'PROCESSO': ['PROCESSO'], 'SERVIDOR': ['SERVIDOR'], 'SIAPE': ['SIAPE'], 'CARGO': ['CARGO'], 'LOTACAO': ['LOTACAO']}
    if 'exclusao' in tipo_gsiste:
        mapeamento_base['EXCLUSAO'] = ['EXCLUSAO']
    return [f"'{c}'" for c, a in mapeamento_base.items() if not any(n in df_columns for n in a)]


# --- Rotas da Aplicação ---
@app.route('/')
def homepage():
    return render_template('index.html')

# === FUNÇÃO ATUALIZADA ===
def handle_individual(creation_function, request_data):
    try:
        nome_ficheiro, doc = creation_function(request_data)
        if not doc:
            logging.error("Falha ao criar documento: Modelo Word não encontrado ou erro de carregamento.")
            return jsonify({"error": "Falha ao criar documento. Verifique se o modelo Word correspondente existe na pasta."}), 500

        # Salvar o documento em um stream de bytes na memória (evita salvar no disco)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0) # Retorna ao início do stream para a leitura

        logging.info(f"Documento individual '{nome_ficheiro}' gerado em memória.")

        # Envia o stream de bytes diretamente como um anexo
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nome_ficheiro,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logging.exception("Erro durante o manuseio da requisição individual:")
        return jsonify({"error": str(e)}), 500

# === FUNÇÃO DE LOTE (JÁ ESTAVA CORRETA) ===
def handle_lote(creation_function, validation_function, request_data, zip_filename):
    if 'arquivo_excel' not in request.files:
        logging.warning("Requisição de lote sem 'arquivo_excel'.")
        return jsonify({"error": "Nenhum ficheiro Excel."}), 400
    try:
        df = pd.read_excel(request.files['arquivo_excel'])
        df.columns = df.columns.str.strip().str.upper()
        logging.info(f"Planilha Excel carregada. Colunas: {df.columns.tolist()}")

        tipo_portaria = None
        if 'TIPO_VACANCIA' in request_data:
            tipo_portaria = request_data.get('TIPO_VACANCIA')
        elif 'TIPO_GSISTE' in request_data:
            tipo_portaria = request_data.get('TIPO_GSISTE')

        if tipo_portaria:
            faltando = validation_function(df.columns, tipo_portaria)
        else:
            faltando = validation_function(df.columns)

        if faltando:
            logging.warning(f"Colunas em falta na planilha: {', '.join(faltando)}")
            return jsonify({"error": "Colunas em falta na planilha: " + ", ".join(faltando)}), 400

        zip_buffer = io.BytesIO()
        portarias_geradas_count = 0
        with ZipFile(zip_buffer, 'w') as zf:
            for index, row in df.iterrows():
                dados_linha = row.to_dict()
                dados_linha.update(request_data)
                try:
                    nome_ficheiro, doc = creation_function(dados_linha)
                    if doc:
                        file_stream = io.BytesIO()
                        doc.save(file_stream)
                        file_stream.seek(0)
                        zf.writestr(nome_ficheiro, file_stream.read())
                        portarias_geradas_count += 1
                        logging.info(f"Portaria '{nome_ficheiro}' adicionada ao ZIP.")
                    else:
                        logging.warning(f"Nenhuma portaria gerada para a linha {index} (dados: {dados_linha}). Verifique o modelo/tipo.")
                except Exception as e_row:
                    logging.error(f"Erro ao gerar portaria para a linha {index}: {e_row}", exc_info=True)

        if portarias_geradas_count == 0:
            logging.warning("Nenhuma portaria pôde ser gerada no lote.")
            return jsonify({"error": "Nenhuma portaria pôde ser gerada. Verifique os dados da planilha e os modelos."}), 400

        zip_buffer.seek(0)
        logging.info(f"Arquivo ZIP '{zip_filename}' gerado com {portarias_geradas_count} portarias.")
        return send_file(zip_buffer, download_name=zip_filename, as_attachment=True, mimetype='application/zip')
    except Exception as e:
        logging.exception("Erro durante o manuseio da requisição em lote:")
        return jsonify({"error": str(e)}), 500


# --- Endpoints (sem alterações na lógica, apenas usam as funções atualizadas) ---
@app.route('/gerar-portaria-gq', methods=['POST'])
def gerar_portaria_gq_endpoint():
    return handle_individual(criar_portaria_gq, request.form.to_dict())

@app.route('/gerar-portaria-gq-lote', methods=['POST'])
def gerar_portaria_gq_lote_endpoint():
    return handle_lote(criar_portaria_gq, validar_colunas_gq, request.form.to_dict(), "portarias_gq.zip")

@app.route('/gerar-portaria-movimentacao', methods=['POST'])
def gerar_portaria_remocao_endpoint():
    return handle_individual(criar_portaria_remocao, request.form.to_dict())

@app.route('/gerar-portaria-movimentacao-lote', methods=['POST'])
def gerar_portaria_remocao_lote_endpoint():
    return handle_lote(criar_portaria_remocao, validar_colunas_remocao, request.form.to_dict(), "portarias_remocao.zip")

@app.route('/gerar-portaria-vacancia', methods=['POST'])
def gerar_portaria_vacancia_endpoint():
    return handle_individual(criar_portaria_vacancia, request.form.to_dict())

@app.route('/gerar-portaria-vacancia-lote', methods=['POST'])
def gerar_portaria_vacancia_lote_endpoint():
    return handle_lote(criar_portaria_vacancia, validar_colunas_vacancia, request.form.to_dict(), "portarias_vacancia.zip")

@app.route('/gerar-portaria-gsiste', methods=['POST'])
def gerar_portaria_gsiste_endpoint():
    return handle_individual(criar_portaria_gsiste, request.form.to_dict())

@app.route('/gerar-portaria-gsiste-lote', methods=['POST'])
def gerar_portaria_gsiste_lote_endpoint():
    return handle_lote(criar_portaria_gsiste, validar_colunas_gsiste, request.form.to_dict(), "portarias_gsiste.zip")


# --- Execução da Aplicação ---
if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=True)